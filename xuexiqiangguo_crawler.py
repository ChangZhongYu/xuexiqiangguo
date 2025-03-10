from bs4 import BeautifulSoup
import re
import os
from docx import Document
from docx.shared import Pt
import time
from playwright.sync_api import sync_playwright
import json
from tools_json import JsonTool
import concurrent.futures
import threading
from tqdm import tqdm

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

# 线程锁，用于保护共享资源
print_lock = threading.Lock()

# 获取标题和正文
def get_article_content(url):
    """
    获取指定URL的文章内容
    
    Args:
        url (str): 文章页面的URL
    
    Returns:
        dict: 包含标题和内容的字典，格式为 {'title': str, 'content': str}
        获取失败时返回None
    """
    try:
        with sync_playwright() as playwright:
            # 初始化浏览器
            browser = playwright.chromium.launch(headless=True)
            context = browser.new_context(user_agent=headers['User-Agent'])
            page = context.new_page()
            
            # 设置请求超时和重试机制
            for attempt in range(3):
                try:
                    page.goto(url, timeout=15000)
                    page.wait_for_selector('div[class*=title]', timeout=10000)
                    # ------------------------等待标签需要更改
                    page.wait_for_selector('div[class*=content]', timeout=10000)
                    break
                except Exception as e:
                    if attempt == 2:
                        raise
                    with print_lock:
                        print(f'请求重试中...（剩余尝试次数：{2 - attempt}）')
                    time.sleep(2)

            # 解析页面内容
            html = page.content()
            soup = BeautifulSoup(html, 'html.parser')

            # 提取文章标题
            title_element = soup.find('div', class_='render-detail-title')
            if not title_element:
                with print_lock:
                    print(f'[错误] 未找到标题元素 | URL: {url}')
                return None
            
            # 提取正文内容
            content_div = soup.find('div', class_="render-detail-content")
            if not content_div:
                with print_lock:
                    print(f'[错误] 未找到正文区域 | URL: {url}')
                return None

            # 处理文本内容
            title = title_element.text.strip()
            paragraphs = [
                p.text.strip() 
                for p in content_div.find_all('p') 
                if p and p.text.strip()
            ]
            
            # 内容有效性验证
            if not paragraphs or len(''.join(paragraphs)) < 50:
                with print_lock:
                    print(f'[警告] 正文内容过短或无效 | 标题: {title}')
                return None

            return {
                'title': title,
                'content': '\n'.join(paragraphs),
                'url': url  # 添加URL以便在多线程环境中识别
            }
            
    except Exception as e:
        with print_lock:
            print(f'[异常] 抓取失败 | URL: {url} | 错误: {str(e)}')
        return None

# 获取所有文章列表网址
def get_article_links(list_url):
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(user_agent=headers['User-Agent'])
            page = context.new_page()
            
            page.goto(list_url)
            html = page.content()
            # print(html)
            try:
                # 定位有效JSON起始位置
                # 使用正则匹配完整JSON数组结构
                json_match = re.search(r'\[.*\]', html, re.DOTALL)
                if not json_match:
                    print('未找到有效JSON数据')
                    return []
                
                # 提取并清理JSON字符串
                json_str = json_match.group(0)
                json_str = re.sub(r'\\\/','/', json_str)  # 处理转义斜杠
                data = json.loads(json_str)
                
                if not isinstance(data, list):
                    print('响应数据结构异常，期望列表')
                    return []
                
                return [item.get('url') for item in data if isinstance(item, dict) and item.get('url')]
            except json.JSONDecodeError as e:
                print(f'JSON解析失败：{str(e)}，响应片段：{json_str[:200]}')
                return []
            except Exception as e:
                print(f'数据处理异常：{str(e)}')
                return []
            
    except Exception as e:
        print(f'获取列表页失败: {str(e)}')
        return []

# 保存文章到Word文档
def save_to_docx(article, timestamp_dir):
    os.makedirs(timestamp_dir, exist_ok=True)
    
    doc = Document()
    
    # 设置文档样式
    styles = doc.styles
    styles['Heading 1'].font.name = '方正小标宋简体'  # 正式公文标题标准字体
    styles['Heading 1'].font.bold = True
    styles['Heading 1'].paragraph_format.alignment = 1  # 居中
    
    normal_style = styles['Normal']
    normal_style.font.name = '仿宋_GB2312'  # 正式公文正文标准字体
    normal_style.font.size = Pt(11)  # 四号字
    normal_style.paragraph_format.line_spacing = 1.5
    
    doc.add_heading(article['title'], level=1)
    for paragraph in article['content'].split('\n'):
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_before = Pt(6)  # 段前间距
    
    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = '生成时间：' + time.strftime('%Y-%m-%d %H:%M:%S')
    
    # 处理文件名
    clean_title = re.sub(r'[\\/:*?\"<>|]', '_', article['title'])[:80]
    filename = f'学习强国文章_{clean_title}.docx'
    
    file_path = os.path.join(timestamp_dir, filename)
    doc.save(file_path)
    with print_lock:
        print(f'文档已保存至: {file_path}')

# 处理单个文章的爬取和保存
def process_article(url, folder, idx, total):
    with print_lock:
        print(f'正在抓取第{idx}/{total}篇: {url}')
    
    article_data = get_article_content(url)
    
    if article_data:
        save_to_docx(article_data, folder)
        return True
    return False

if __name__ == '__main__':
    # 用户交互模块
    print('='*40)
    print('学习强国文章抓取工具 (多线程版)')
    print('='*40)
    
    # 设置线程数
    max_workers = input('请输入线程数 (默认为4): ')
    max_workers = int(max_workers) if max_workers.isdigit() and int(max_workers) > 0 else 4
    print(f'已设置线程数: {max_workers}')
    
    field = input('请输入指定字段及专题名称英文逗号隔开：')
    if not field:
        # 读取JSON文件
        json_data = JsonTool.read_json_file('resource\学习强国专栏ID.json')
        subject = input("需要下载的主题关键字：")
        specific_values = JsonTool.extract_specific_values(json_data, subject)
        if not specific_values:
            print('未读取到有效主题，程序退出')
            exit()
    else:
        # 校验field格式
        if len(field.split(',')) != 2:
            print('输入格式错误，请使用正确的格式')
            exit()
        specific_values = {field.split(',')[0]: field.split(',')[1]}
    
    # 测试标签，用于打断循环
    flag = 0
    for item, folder in specific_values.items():
        # 测试标签
        # if flag > 2:
        #     break
        # 拼接文章列表URL
        list_url = f"https://www.xuexi.cn/lgdata/{item}.json"
        print(f"获取文章列表: {list_url} -> {folder}")
        
        # 获取文章列表链接
        article_urls = get_article_links(list_url)
        total_articles = len(article_urls)
        print(f'发现{total_articles}篇文章')
        
        if not article_urls:
            continue
        
        # 创建进度条
        progress_bar = tqdm(total=total_articles, desc=f"下载进度", unit="篇")
        
        # 使用线程池并行处理文章
        successful_count = 0
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # 提交所有任务
            future_to_url = {executor.submit(process_article, url, folder, idx+1, total_articles): url 
                             for idx, url in enumerate(article_urls)}
            
            # 处理完成的任务
            for future in concurrent.futures.as_completed(future_to_url):
                url = future_to_url[future]
                try:
                    if future.result():
                        successful_count += 1
                except Exception as exc:
                    with print_lock:
                        print(f'处理文章时发生异常: {url} - {exc}')
                finally:
                    progress_bar.update(1)
        
        progress_bar.close()
        print(f'专题 "{folder}" 处理完成，成功下载 {successful_count}/{total_articles} 篇文章')
        
        # 添加专题间的间隔
        time.sleep(2)
    
    print('\n所有专题处理完成!')